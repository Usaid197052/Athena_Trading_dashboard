"""Generate Athena Trading multi-agent architecture/design DOCX.

Run from the repo root:
    python docs/generate_architecture_docx.py
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent / "Athena_Trading_Multi_Agent_Architecture.docx"

NAVY = RGBColor(0x1A, 0x14, 0x10)
BRONZE = RGBColor(0xC1, 0x50, 0x2E)
MUTED = RGBColor(0x5A, 0x4A, 0x3A)


def _set_run(run, *, size=11, bold=False, color=NAVY, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        _set_run(run, size={1: 18, 2: 14, 3: 12}.get(level, 11), bold=True, color=BRONZE if level == 1 else NAVY)
    return p


def para(doc, text, *, bold=False, italic=False, size=11, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    _set_run(run, size=size, bold=bold, italic=italic)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            _set_run(run, size=11)


def cpm(doc, current: str, proposed: str, migration: str):
    add_heading(doc, "Current State", 3)
    para(doc, current)
    add_heading(doc, "Proposed State", 3)
    para(doc, proposed)
    add_heading(doc, "Migration Path", 3)
    para(doc, migration)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _set_run(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shading = cell._tePr if False else cell._tc.get_or_add_tcPr()
        # simple text only; skip complex shading for compatibility
    for r_i, row in enumerate(rows, start=1):
        for c_i, val in enumerate(row):
            cell = t.rows[r_i].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            _set_run(run, size=10)
    doc.add_paragraph()
    return t


def build() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("ATHENA")
    _set_run(r, size=28, bold=True, color=BRONZE)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run("Trading Dashboard — Multi-Agent Architecture & Algorithm Design")
    _set_run(r, size=16, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        f"Design document  ·  {date.today().isoformat()}  ·  "
        "For developers and the project owner"
    )
    _set_run(r, size=11, italic=True, color=MUTED)

    para(
        doc,
        "This document describes how the existing Athena Trading Desk becomes a local "
        "multi-agent analysis system. It is not financial advice. The new agents never "
        "place orders. Existing demo auto-trade remains a separate, unchanged path.",
        italic=True,
    )
    para(
        doc,
        "How to read this document: every major chapter uses the same three headings — "
        "Current State, Proposed State, and Migration Path — so you can see what already "
        "works, what we are adding, and how we get there without breaking the desk.",
    )

    add_heading(doc, "Locked decisions", 1)
    bullets(doc, [
        "Athena the orchestrator uses Gemini Flash (existing API), text only — no microphone and no spoken replies.",
        "DeepSeek-R1 7B (quantized) is the fundamental / news agent.",
        "Qwen2.5-Coder 7B (quantized) is the technical / algorithm agent.",
        "A third graph/numerical agent is optional, disabled, and not installed in this version.",
        "Local models run one at a time on the laptop GPU (sequential load/unload via Ollama).",
        "Existing MetaTrader 5 demo auto-trade stays as it is. Agent output is analysis only and never calls order_send.",
        "Speech-to-text is removed from the Trading Dashboard only. Personal-assistant STT is not in this repository.",
        "API keys are stored locally on Windows using DPAPI encryption, never in git or in logs.",
    ])

    # 1
    add_heading(doc, "1. Current architecture assessment", 1)
    cpm(
        doc,
        "Athena is a Windows desktop app (Python 3.12, PyQt6). trading_main.py starts a Gemini Live "
        "voice session. The model talks, calls tools (trading_desk, mt5_analysis, trading_control, "
        "web_search), and a 20-second watch-loop may place demo orders through MetaTrader 5. "
        "Technical numbers (EMA, RSI, MACD, ATR, Bollinger, support/resistance) are already computed "
        "in NumPy inside mt5_analysis.py. News is a short DuckDuckGo fetch. There is no local LLM, "
        "no Ollama, no pytest suite, and API keys sit in plaintext config/api_keys.json (gitignored).",
        "Keep the HUD, MT5 connection, deterministic indicators, demo executor, and watch-loop. "
        "Add a parallel analysis pipeline: data adapters → local worker agents → Gemini Flash "
        "aggregation → dashboard. Voice Live is replaced by typed commands plus an Analyze button.",
        "Additive packages (agents/, data/, security/, tests/). Do not rewrite ui.py. Do not change "
        "place_market behaviour. Replace only the Gemini Live audio loop in trading_main.py.",
    )
    bullets(doc, [
        "Entry point: python trading_main.py (unchanged).",
        "HUD: AthenaUI(trading_mode=True) — Pause / Resume / Flatten already exist.",
        "Personal assistant (main.py) is a separate project and is not in this tree.",
        "Leftover WhatsApp / NYX code stays unused; it is not merged into the agent system.",
    ])

    # 2
    add_heading(doc, "2. Proposed multi-agent architecture", 1)
    cpm(
        doc,
        "One cloud model (Gemini Live) is the entire brain: speech, tools, briefing, and implicit FA.",
        "Athena (Gemini Flash) is the project handler. It collects a market snapshot, runs "
        "deterministic indicators, then sequentially calls Agent 2 (Qwen) and Agent 1 (DeepSeek), "
        "optionally Agent 3, validates structured JSON, resolves conflicts in Python, and shows a "
        "human-readable assessment. Auto-trade continues on the old scorecard and never reads agent bias.",
        "Introduce TradingSession in trading_main.py. Wire Analyze and typed 'analyze' to "
        "agents.orchestrator.run_market_analysis. Keep watch_tick → trading_desk isolated.",
    )
    para(doc, "Flow (plain language): you ask for analysis → Athena gathers prices and news → "
         "the built-in calculator produces indicator numbers → Qwen explains the chart numbers → "
         "DeepSeek reads the news → Athena combines both and tells you where they agree or disagree.")
    para(doc, "Technical flow:", italic=True)
    para(
        doc,
        "User / Analyze → Orchestrator → Data acquisition + IndicatorEngine → "
        "Qwen (GPU) → unload → DeepSeek (GPU) → unload → optional Graph → "
        "conflict resolver + Flash narrative → HUD. Parallel: watch_tick → trading_desk → mt5_executor (demo).",
    )

    # 3
    add_heading(doc, "3. Agent responsibilities", 1)
    cpm(
        doc,
        "No worker agents. Gemini Live both chats and calls trading_desk, which mixes TA, a little FA, gates, and execution.",
        "Agent 1 (DeepSeek-R1 7B): classify news (material / low-impact / noise), implications, risks, facts vs interpretation. "
        "No price targets. Agent 2 (Qwen2.5-Coder 7B): interpret the existing indicator snapshot; optional algorithm drafts that never run until tests pass. "
        "Agent 3 (optional): numerical market-structure comments on OHLC tables — not screenshots. "
        "Athena: dispatch, validate, conflict resolution, presentation.",
        "New modules under agents/. Prompts require JSON. Invalid JSON is rejected and Athena continues with whatever else succeeded.",
    )

    # 4
    add_heading(doc, "4. Model selection", 1)
    cpm(
        doc,
        "Gemini Live native-audio for the desk; Gemini 3.6 Flash for search; Flash Lite stored but unused on the trading path. No local models.",
        "Orchestrator: Gemini Flash (user-selectable ID). Worker 1: Ollama tag deepseek-r1:7b (Q4_K_M). "
        "Worker 2: qwen2.5-coder:7b (Q4_K_M). Worker 3 (later, not installed): qwen2.5-math:7b, with phi4-mini as a lighter fallback.",
        "User installs Ollama and pulls the two 7B tags when ready. Athena does not download models at import time.",
    )
    table(
        doc,
        ["Role", "Model", "Why"],
        [
            ["Orchestrator", "Gemini Flash", "Already in the app; no VRAM; good at combining structured findings"],
            ["Fundamentals", "DeepSeek-R1 7B Q4", "Reasoning model; news implications; mandatory"],
            ["Technical", "Qwen2.5-Coder 7B Q4", "Reads numbers and can draft algorithms; mandatory"],
            ["Graph (later)", "Qwen2.5-Math 7B Q4", "Numerical structure; optional plug-in"],
        ],
    )

    # 5
    add_heading(doc, "5. Quantization strategy", 1)
    cpm(
        doc,
        "Not applicable — inference is cloud Gemini.",
        "Use Ollama default Q4_K_M. Context cap 4096 tokens. num_predict capped so DeepSeek thinking traces cannot fill VRAM. "
        "Do not use Q8 or 14B/32B on this laptop. Keep-alive 0 so weights leave the GPU after each agent.",
        "Record model tags and options in config/agents.json. Tune num_predict after the first measured run on this machine.",
    )

    # 6
    add_heading(doc, "6. Hardware feasibility", 1)
    para(
        doc,
        "Target laptop: Intel Core i9-13980HX, NVIDIA RTX 4070 Laptop GPU, about 15.6 GB usable RAM, Windows. "
        "The 4070 Laptop ships with 8 GB GDDR6 in standard SKUs. nvidia-smi was not readable without administrator rights during planning; "
        "implementation should confirm VRAM once. Sequential design remains correct even if VRAM is 12 GB.",
    )
    cpm(
        doc,
        "GPU is used only for HUD utilisation meters. All LLM work is cloud-side. MT5, PyQt, and Python share RAM.",
        "One 7B Q4_K_M model needs roughly 5.8–6.5 GB VRAM at 4k context. Two at once need ~12 GB and will not fit. "
        "Keeping a second 7B on CPU is not practical with ~16 GB system RAM plus Windows, MT5, and the HUD. "
        "Gemini Flash adds almost no VRAM.",
        "Never load two local models together. Unload after each worker. If GPU OOM occurs, retry once with a smaller num_predict; then mark the agent unavailable.",
    )
    table(
        doc,
        ["Combination", "Practical?"],
        [
            ["DeepSeek 7B Q4 on GPU alone", "Yes"],
            ["Qwen 7B Q4 on GPU alone", "Yes"],
            ["Both 7Bs on GPU together", "No"],
            ["One 7B GPU + one 7B CPU", "No (RAM)"],
            ["Gemini Flash + one 7B + Athena + MT5", "Yes"],
            ["Third 7B while another is loaded", "No"],
        ],
    )

    # 7
    add_heading(doc, "7. Sequential vs simultaneous execution", 1)
    cpm(
        doc,
        "Gemini Live is always connected while the desk is awake. Tools run on a thread pool. Watch-loop every 20 seconds.",
        "CPU work first (MT5 bars, indicators, news). Then load Qwen, interpret TA, unload. Then load DeepSeek, interpret news, unload. "
        "Then optional graph agent. Then Flash writes the combined explanation. Watch-loop is independent and does not load local models.",
        "Ollama HTTP keep_alive=0. Orchestrator owns the lock so two analyses cannot overlap and fight for VRAM.",
    )
    para(doc, "Order: data → indicators (HUD can already show numbers) → Agent 2 → Agent 1 → Agent 3 if enabled → aggregation.")
    para(doc, "Timeouts: Qwen 90s, DeepSeek 180s (reasoning), graph 90s, Flash 45s. Timeout = that layer is missing, not a crash.")

    # 8
    add_heading(doc, "8. Dependency analysis", 1)
    cpm(
        doc,
        "requirements.txt: PyQt6, sounddevice, google-genai, numpy, pandas, MetaTrader5, pywin32, psutil, ddgs, vosk, opencv, mss, requests, beautifulsoup4. No yfinance, pytest, pydantic, Ollama.",
        "Keep existing packages. Add pydantic (structured agent JSON) and pytest (+ timeout plugin) for tests. Talk to Ollama over HTTP with requests (no extra inference framework). "
        "python-docx is only for regenerating this design file. yfinance is an optional fallback adapter, off by default. "
        "Do not add pandas-ta, backtrader, or vectorbt — they overlap the NumPy engine and add maintenance cost.",
        "Update requirements.txt. Vosk and sounddevice remain listed so leftover assistant files still import; the trading process simply does not start them.",
    )

    # 9
    add_heading(doc, "9. Market-data architecture", 1)
    cpm(
        doc,
        "mt5_analysis.py talks to MetaTrader 5 directly. trading_desk and Gemini tools call it. News is _news_fast (2.5s, four headlines) plus optional calendar.",
        "A data layer sits between the world and the agents: MT5 adapter (bars, ticks, spread, calendar), news adapter (DuckDuckGo with a longer timeout, optional Flash grounding), "
        "optional yfinance adapter if MT5 does not have the symbol. Output is NormalizedMarketSnapshot. Agents never import MetaTrader5 or yfinance.",
        "Wrap existing functions; do not duplicate indicator math. Keep the 2.5s news path for auto-trade blackout gates. Use the longer news path only for Agent 1.",
    )

    # 10
    add_heading(doc, "10. MetaTrader 5 integration", 1)
    cpm(
        doc,
        "Direct MetaTrader5 package IPC: attach to running terminal64.exe, copy_rates_from_pos (260 bars), ticks, positions, demo-only order_send, JPEG window snapshot for Live vision.",
        "Same IPC, behind Mt5MarketAdapter. Snapshot/vision is not used by the new analysis pipeline (structured numbers only). Executor stays demo-only, magic 20260820, mandatory SL/TP.",
        "Do not spread mt5.initialize through agents. Only actions/mt5_*.py may import MetaTrader5. Adapter returns dicts/lists.",
    )

    # 11
    add_heading(doc, "11. API-key architecture", 1)
    cpm(
        doc,
        "Single gemini_api_key string in gitignored JSON. Settings overlay can replace it and pick Live / Flash / Lite IDs. Mask helper exists. No encryption. No test button. No multiple keys.",
        "Multiple labelled keys. Windows DPAPI (CryptProtectData) via pywin32. Enable/disable, set active, test with a tiny Flash ping, remove. After storage the UI shows only a mask (first 8 … last 4). "
        "Trading-mode settings hide Live and Lite dropdowns. Flash dropdown stays and remains data-driven so future model IDs do not need a UI rewrite. Keys never appear in logs, prompts, or git.",
        "On first load, if a plaintext gemini_api_key is present, encrypt it, assign an id, mark it active, rewrite the file. get_gemini_key() decrypts the active enabled key.",
    )

    # 12
    add_heading(doc, "12. Logging architecture", 1)
    cpm(
        doc,
        "logs/trading/desk.log, decisions.jsonl, analysis.log, analysis.jsonl, logs/mt5/*. Developer-oriented lines such as DESK start EURUSD H1. Optional HUD sink.",
        "Keep those files. Add logs/trading/activity.log — sentences a non-technical owner can read (DeepSeek has finished reviewing today's market news). "
        "Add logs/trading/agents.debug.log plus JSONL for timings, model tags, validation, load/unload, redacted HTTP errors. Both streams refuse to print API keys.",
        "core/activity_logger.py and core/agent_debug_logger.py. HUD continues to use write_log for activity lines.",
    )

    # 13
    add_heading(doc, "13. Security architecture", 1)
    cpm(
        doc,
        "Gitignore covers api_keys.json, logs, models/. Remote password uses PBKDF2. Gemini never receives order_send. Live account orders are blocked. No sandbox because no generated code runs.",
        "DPAPI for Gemini keys. Log sanitizer for AIza and Bearer tokens. Agent 2 default path does not execute code. Algorithm drafts go to sandbox/jobs/<id>/ and run in a subprocess with a tiny allowed-import list (numpy, pandas, math), no network, timeout, no MetaTrader5. "
        "Promotion to sandbox/approved/ only after golden tests. Agents never receive the API key in prompts.",
        "security/keystore.py and security/sanitize.py. sandbox/runner.py. gitignore sandbox/jobs.",
    )

    # 14
    add_heading(doc, "14. Testing architecture", 1)
    cpm(
        doc,
        "No tests directory, no pytest.",
        "Unit tests for indicators, scoring, symbol aliases, AgentResult schema, conflict table, sanitizer, keystore (Windows). "
        "Integration tests with fake Ollama/Gemini/MT5. Optional slow tests that ping a real Ollama if present. Algorithm gate: generated SMA must match NumPy or be rejected.",
        "Add tests/ and pytest.ini. Default CI/dev run must not require GPU or live brokers.",
    )

    # 15
    add_heading(doc, "15. Trading-analysis algorithm", 1)
    cpm(
        doc,
        "A numeric score in mt5_analysis._score (EMA stack, RSI vs stack, MACD vs stack, proximity to S/R, candle pattern) maps to BUY / SELL / WAIT. FA is mostly a news blackout, not a reasoned bias. That score also drives auto-trade.",
        "Two layers. Layer A (unchanged): desk scorecard → auto-trade gates → optional demo order. Layer B (new, display only): "
        "Fundamental bias + Technical bias + optional Market structure + risk + confidence → Overall assessment. "
        "Outcomes: Strong Bullish, Bullish, Neutral, Bearish, Strong Bearish, Insufficient Data. Never average two confidence percentages.",
        "Implement conflict_resolve() in Python. Flash may only paraphrase the structured result. Do not feed Layer B into place_market.",
    )
    add_heading(doc, "Conflict rules (plain language)", 2)
    bullets(doc, [
        "If both agents fail: Insufficient Data.",
        "If only one side arrives: show it, say the other side is missing, never call it Strong.",
        "If they agree and both look solid (including a clear technical score): Strong Bullish or Strong Bearish.",
        "If they agree more weakly: Bullish or Bearish.",
        "If they disagree: never Strong. Fresh high-impact news can lean fundamental one notch, but the screen must say they disagree and recommend waiting for confirmation. A very strong chart against junk news can lean technical, still with a disagreement flag. Otherwise Neutral / high uncertainty.",
        "Always list risks from both sides.",
    ])

    # 16
    add_heading(doc, "16. Agent communication protocol", 1)
    cpm(
        doc,
        "Tools return free-form text cards (BIAS BUY, WHY, PLAN, GATES) meant to be spoken.",
        "Every worker returns a JSON object validated by Pydantic: agent, task, input_data, analysis, signals, confidence, reasoning_summary, "
        "risks, warnings, timestamp, model, model_version, execution_status, plus bias, facts, interpretations, scenarios, data_quality, recency_sec. "
        "The HUD explanation is human-readable; the contract between processes is the object.",
        "Ollama requests use format=json. Strip DeepSeek <think> traces before parse. One repair retry, then rejected.",
    )

    # 17
    add_heading(doc, "17. Conflict-resolution mechanism", 1)
    cpm(
        doc,
        "No second opinion. The desk score is the only bias.",
        "Python function, not an LLM vote. Inputs: each agent's bias and confidence, FA event severity and recency, TA engine score and whether Qwen agrees with the engine, data quality, missing flags. "
        "Output includes disagreement: true/false, why, and recommendation. Historical reliability can be a later weight; v1 does not pretend we have a live track record.",
        "Table-driven unit tests cover agree, disagree, missing agent, high-impact FA vs weak TA, extreme TA vs noise FA.",
    )

    # 18
    add_heading(doc, "18. Failure and recovery strategy", 1)
    cpm(
        doc,
        "Gemini Live reconnects on drop. Invalid API key re-opens setup. MT5 keepalive recovers IPC. News timeout skips headlines. Live account blocks orders.",
        "Same MT5/demo behaviour. New: Ollama down → workers unavailable, Flash can still say so. GPU OOM → unload, mark error, continue. Malformed JSON → rejected. "
        "yfinance unused unless enabled and MT5 missed the symbol. Agent disagreement is not a failure — it is a first-class outcome. "
        "Partial analysis is shown: 'Fundamental analysis is unavailable. This result is based only on technical data.'",
        "Every agent returns a status enum. Orchestrator never raises to the UI thread; it always produces an assessment object.",
    )

    # 19
    add_heading(doc, "19. UI changes", 1)
    cpm(
        doc,
        "Content panel under the owl shows account, MT5, auto-trade, positions, LAST why/plan/gates. Pause Resume Flatten. Voice listening state. Settings: one key plus three model combos including Live voice.",
        "Same panel, plus agent status dots (Athena, DeepSeek, Qwen, Graph disabled), an Analyze button, human activity lines in the log, and an ANALYSIS block separate from LAST (so auto-trade and agent views are not mixed). "
        "Settings in trading mode: key list, add/remove/enable/activate/test, Flash dropdown only. Typed commands remain the main way to talk to Athena.",
        "Extend ui.py in trading_mode only. Do not rebuild the HUD. Hide Live/Lite combos when trading_mode is true.",
    )

    # 20
    add_heading(doc, "20. Personal-assistant separation", 1)
    cpm(
        doc,
        "This repository is already trading-only. readme.md says the personal assistant is another project. Leftover WhatsApp modules and NYX overlay branches still exist in ui.py, gated by trading_mode.",
        "Keep that split. Do not import WhatsApp or remote-phone dashboard into the agent manager. Do not delete wakeword.py in this pass; simply stop calling it from trading_main.py. Sleep/wake uses the tray icon.",
        "No mega-refactor. TradingSession replaces TradingLive. Assistant leftovers stay dormant.",
    )

    # 21
    add_heading(doc, "21. Proposed project structure", 1)
    para(doc, "Existing files stay. New folders are additive:")
    bullets(doc, [
        "agents/ — protocol, Ollama runtime, manager, fundamental, technical, graph stub, orchestrator, sandbox helper, prompts, state",
        "data/ — snapshot model; adapters for MT5, news, optional yfinance",
        "security/ — DPAPI keystore, log sanitizer",
        "core/activity_logger.py, core/agent_debug_logger.py",
        "sandbox/runner.py — approved import whitelist",
        "config/agents.json — models, timeouts, graph enabled flag",
        "tests/ — unit and integration",
        "docs/ — this design document and its generator",
    ])
    para(doc, "trading_main.py, ui.py, actions/trading_desk.py, actions/mt5_analysis.py, actions/mt5_executor.py remain the operational core of the desk.")

    # 22
    add_heading(doc, "22. Implementation phases", 1)
    cpm(
        doc,
        "A working Gemini Live demo desk with auto-trade.",
        "Phased delivery that always leaves auto-trade working even if Ollama is not installed yet.",
        "0 design DOCX (this file). 1 protocol + Ollama + logs. 2 data adapters. 3 indicator tests. 4 DeepSeek agent. 5 Qwen agent + sandbox scaffold. "
        "6 Flash aggregation + conflicts. 7 remove Live STT from trading_main. 8 UI status/Analyze. 9 encrypted multi-key. 10 graph stub. 11–14 sandbox tests, security, pytest, latency caps.",
    )

    # 23
    add_heading(doc, "23. Risks and mitigations", 1)
    table(
        doc,
        ["Risk", "Mitigation"],
        [
            ["8 GB VRAM cannot hold two 7Bs", "Sequential unload; analysis lock"],
            ["DeepSeek thinking is slow / large", "Timeouts, num_predict cap, strip <think>"],
            ["Ollama not installed", "Degrade to indicators + Flash message; desk still trades on old scorecard"],
            ["Gemini quota / key failure", "Show structured assessment without narrative; prompt to test/rotate keys"],
            ["Qwen emits executable Python", "Default path interpret-only; sandbox whitelist if code is requested later"],
            ["Agents treated as trading signals", "HUD labels analysis as decision support; no order_send from agents"],
            ["News quality / DDG blocks", "Longer timeout, optional Flash grounding, calendar + static high-impact windows"],
            ["MT5 down", "Snapshot fails closed; assessment Insufficient Data; keepalive unchanged"],
            ["Key leakage in logs", "Sanitizer + never pass keys into prompts"],
            ["STT accidentally removed from assistant", "Assistant is another repo; this change is trading_main only"],
        ],
    )

    # 24
    add_heading(doc, "24. Future extension points", 1)
    bullets(doc, [
        "Enable Agent 3 (qwen2.5-math:7b) behind config/agents.json without changing the orchestrator contract.",
        "Optional TTS later (not in this version).",
        "Reliability weights once a journal of agent vs subsequent price exists.",
        "More indicators (ADX, Stochastic, VWAP) as tested NumPy plugins promoted through the sandbox.",
        "yfinance fallback for cash equities that MT5 does not list.",
        "User-approved wiring of consensus into demo auto-trade — only if explicitly requested later. Default remains isolated.",
        "Swap Ollama tags without UI rewrite (config + protocol model field).",
    ])

    add_heading(doc, "What this system will not do", 1)
    bullets(doc, [
        "It will not place live-account orders (existing hard rule).",
        "It will not treat an agent percentage as a guaranteed forecast.",
        "It will not run untested generated code on your machine with full permissions.",
        "It will not download multi-gigabyte models unless you install Ollama and pull them.",
        "It will not merge the personal assistant into the trading desk.",
    ])

    add_heading(doc, "Owner summary (non-technical)", 1)
    para(
        doc,
        "Athena on the trading screen becomes a coordinator. One specialist reads news. Another reads the numbers MetaTrader already computed. "
        "Athena tells you whether they agree. The old automatic demo trader keeps using the simple scorecard it uses today, unless you later ask to change that. "
        "You type, or press Analyze — the microphone on this dashboard is turned off. Your Gemini keys stay on this PC, encrypted by Windows, and the full key is not shown again after you save it.",
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("End of design document  ·  Implementation must follow Current → Proposed → Migration in each area.")
    _set_run(r, size=10, italic=True, color=MUTED)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")
